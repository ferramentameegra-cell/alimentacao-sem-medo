#!/usr/bin/env python3
"""
Extrai base de conhecimento dos arquivos .docx do Planeta Intestino.
FONTE ÚNICA: data/pdfs/*.docx (NUNCA usar cardapios-planeta-intestino.pdf)
"""

import json
import re
from pathlib import Path
from docx import Document
from typing import List, Dict, Any, Optional

# Mapeamento: nome do arquivo (sem extensão) -> condicao_digestiva
MAPEAMENTO_CONDICAO = {
    'Azia e Refluxo': 'azia_refluxo',
    'Bloqueio Defecatório': 'intestino_preso',
    'Colite': 'colite',
    'Dieta Anti-inflamatória': 'anti_inflamatoria',
    'Disbiose': 'disbiose',
    'Diverticulite': 'diverticulite',
    'Divertículos_': 'diverticulos_intestinais',
    'Gases': 'gases_abdome_distendido',
    'INTESTINO PRESO': 'intestino_preso',
    'Intolerancia à Lactose': 'intolerancia_lactose',
    'Má Digestão': 'ma_digestao',
    'Prevenção a diarreia': 'diarreia',
    'sem gluten e lactose': 'sem_gluten_lactose',
    'Sem Gluten': 'sem_gluten',
    'SII': 'sindrome_intestino_irritavel',
    'zJantar casual_romantico': 'geral',
}

# Palavras-chave para identificar tipo de refeição no texto
PALAVRAS_REFEICAO = {
    'cafe_manha': ['café da manhã', 'cafe da manha', 'café da manhá', 'desjejum', 'manhã'],
    'almoco': ['almoço', 'almoco', 'almoço'],
    'lanche_tarde': ['lanche', 'lanche da tarde', 'tarde', 'merenda'],
    'jantar': ['jantar', 'ceia', 'jantar casual'],
}

def extrair_texto_docx(caminho: Path) -> str:
    """Extrai todo o texto de um arquivo .docx"""
    doc = Document(caminho)
    partes = []
    for para in doc.paragraphs:
        if para.text.strip():
            partes.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    partes.append(cell.text)
    return '\n'.join(partes)

def identificar_tipo_refeicao(texto_antes: str) -> Optional[str]:
    """Identifica tipo de refeição pelo contexto do texto anterior"""
    texto_lower = texto_antes.lower()
    for tipo, palavras in PALAVRAS_REFEICAO.items():
        for p in palavras:
            if p in texto_lower:
                return tipo
    return None

def extrair_itens_texto(texto: str, condicao: str, nome_arquivo: str) -> List[Dict[str, Any]]:
    """Extrai itens alimentares (nome + quantidade) do texto"""
    itens = []
    linhas = texto.split('\n')
    tipo_atual = 'almoco'  # default
    
    # Padrões para quantidade: 100g, 200ml, 2 colheres, 1 xícara, 1 unidade, etc.
    padrao_quant = re.compile(
        r'(\d+[,.]?\d*)\s*'
        r'(g|kg|ml|l|colher|colheres|xícara|xícaras|fatia|fatias|unidade|unidades|'
        r'prato|porção|porções|gramas|ml|copo|copos|colher de sopa|colher de chá)\b',
        re.IGNORECASE
    )
    
    # Padrão alternativo: quantidade no final da linha
    padrao_quant_final = re.compile(r'[-–—]\s*(\d+[,.]?\d*\s*(?:g|kg|ml|l|colher|colheres|xícara|fatia|unidade|prato|porção|gramas|copo)s?)\s*$', re.IGNORECASE)
    
    for i, linha in enumerate(linhas):
        linha = linha.strip()
        if len(linha) < 4:
            continue
        
        # Verificar se é cabeçalho de seção
        tipo_detectado = identificar_tipo_refeicao(linha)
        if tipo_detectado and len(linha) < 80:
            tipo_atual = tipo_detectado
            continue
        
        # Tentar extrair item com quantidade
        match = padrao_quant.search(linha)
        if match:
            qtd_inicio = match.start()
            nome = linha[:qtd_inicio].strip()
            quantidade = match.group(0).strip()
            # Limpar nome
            nome = re.sub(r'^[-•\d.)\s]+', '', nome).strip()
            nome = re.sub(r'\s+', ' ', nome)
            if nome and len(nome) > 2 and len(nome) < 100:
                itens.append({
                    'nome': nome,
                    'quantidade': quantidade,
                    'tipo': tipo_atual,
                    'condicao_digestiva': condicao,
                    'fonte': nome_arquivo,
                })
            continue
        
        # Padrão " - 100g" no final
        match_final = padrao_quant_final.search(linha)
        if match_final:
            nome = linha[:match_final.start()].strip()
            quantidade = match_final.group(1).strip()
            nome = re.sub(r'^[-•\d.)\s]+', '', nome).strip()
            if nome and len(nome) > 2:
                itens.append({
                    'nome': nome,
                    'quantidade': quantidade,
                    'tipo': tipo_atual,
                    'condicao_digestiva': condicao,
                    'fonte': nome_arquivo,
                })
    
    return itens

def processar_docx(caminho: Path) -> List[Dict[str, Any]]:
    """Processa um arquivo .docx e retorna lista de itens"""
    nome_base = caminho.stem
    condicao = MAPEAMENTO_CONDICAO.get(nome_base, 'geral')
    
    try:
        texto = extrair_texto_docx(caminho)
        itens = extrair_itens_texto(texto, condicao, nome_base + '.docx')
        return itens
    except Exception as e:
        print(f"  ⚠ Erro ao processar {caminho.name}: {e}")
        return []

def main():
    base_dir = Path(__file__).parent.parent
    pdfs_dir = base_dir / 'data' / 'pdfs'
    output_path = base_dir / 'data' / 'base_conhecimento.json'
    
    # Arquivos permitidos (apenas .docx, NUNCA o PDF)
    arquivos_permitidos = [
        'Azia e Refluxo.docx',
        'Bloqueio Defecatório.docx',
        'Colite.docx',
        'Dieta Anti-inflamatória.docx',
        'Disbiose.docx',
        'Diverticulite.docx',
        'Divertículos_.docx',
        'Gases.docx',
        'INTESTINO PRESO.docx',
        'Intolerancia à Lactose.docx',
        'Má Digestão.docx',
        'Prevenção a diarreia.docx',
        'sem gluten e lactose.docx',
        'Sem Gluten.docx',
        'SII.docx',
        'zJantar casual_romantico.docx',
    ]
    
    print("📚 Extraindo base de conhecimento dos arquivos .docx")
    print("   Fonte: data/pdfs/*.docx (PDF excluído permanentemente)\n")
    
    todos_itens = []
    contador_ids = 0
    
    for nome_arquivo in arquivos_permitidos:
        caminho = pdfs_dir / nome_arquivo
        if not caminho.exists():
            print(f"  ⚠ Não encontrado: {nome_arquivo}")
            continue
        
        itens = processar_docx(caminho)
        for item in itens:
            contador_ids += 1
            item['id'] = f"docx_{contador_ids:04d}"
            todos_itens.append(item)
        
        print(f"  ✓ {nome_arquivo}: {len(itens)} itens")
    
    # Deduplicar por (nome, quantidade, tipo, condicao) mantendo primeira ocorrência
    vistos = set()
    itens_unicos = []
    for item in todos_itens:
        chave = (item['nome'].lower(), item['quantidade'], item['tipo'], item['condicao_digestiva'])
        if chave not in vistos:
            vistos.add(chave)
            itens_unicos.append(item)
    
    resultado = {
        'itens': itens_unicos,
        'total_itens': len(itens_unicos),
        'fontes': arquivos_permitidos,
        'origem': 'Arquivos .docx do Planeta Intestino (PDF excluído)',
    }
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(resultado, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ Base salva em: {output_path}")
    print(f"   Total de itens: {len(itens_unicos)}")
    
    # Amostra
    if itens_unicos:
        print("\n📄 Amostra (5 primeiros):")
        for item in itens_unicos[:5]:
            print(f"   - {item['nome']} — {item['quantidade']} ({item['tipo']}, {item['condicao_digestiva']})")

if __name__ == '__main__':
    main()
